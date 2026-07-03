"""DOCX conversion handler."""


class DocxConverter:
    """Convert DOCX to Markdown, HTML, and text."""

    def convert(self, input_path, target_format, output_path):
        if target_format == "md":
            self._to_markdown(input_path, output_path)
        elif target_format == "html":
            self._to_html(input_path, output_path)
        elif target_format == "txt":
            self._to_text(input_path, output_path)
        else:
            raise ValueError(f"Unsupported DOCX target format: {target_format}")

    def _read_docx(self, input_path):
        from docx import Document
        return Document(input_path)

    def _paragraph_to_md(self, paragraph):
        text = paragraph.text.strip()
        if not text:
            return ""

        # Determine heading level by style name (Heading 1, Heading 2, ...)
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
                return "#" * level + " " + text
            except (ValueError, IndexError):
                pass

        # Basic inline formatting
        md = ""
        for run in paragraph.runs:
            run_text = run.text
            if run.bold:
                run_text = f"**{run_text}**"
            if run.italic:
                run_text = f"*{run_text}*"
            md += run_text
        return md.strip()

    def _table_to_md(self, table):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        lines = []
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
        for row in rows[1:]:
            # Pad row to match header length
            while len(row) < len(rows[0]):
                row.append("")
            lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")
        return "\n".join(lines)

    def _to_markdown(self, input_path, output_path):
        doc = self._read_docx(input_path)
        parts = []

        for block in doc.inline_shapes:
            # Images are not directly iterable here; we scan paragraphs instead.
            pass

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]
            if tag == "p":
                # Find corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element is element:
                        md = self._paragraph_to_md(para)
                        if md:
                            parts.append(md)
                        break
            elif tag == "tbl":
                for table in doc.tables:
                    if table._element is element:
                        parts.append(self._table_to_md(table))
                        parts.append("")
                        break
            elif tag == "sectPr":
                continue

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(parts))

    def _to_html(self, input_path, output_path):
        md_path = output_path + ".tmp.md"
        try:
            self._to_markdown(input_path, md_path)
            with open(md_path, "r", encoding="utf-8") as f:
                md = f.read()
            try:
                import markdown
                html = markdown.markdown(md, extensions=["tables"])
            except ImportError:
                # Fallback: wrap markdown in pre if markdown package not installed
                html = "<pre>" + md.replace("&", "&amp;").replace("<", "&lt;") + "</pre>"
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(f"<!DOCTYPE html>\n<html><head><meta charset=\"utf-8\"></head><body>\n{html}\n</body></html>")
        finally:
            if __import__("os").path.exists(md_path):
                __import__("os").unlink(md_path)

    def _to_text(self, input_path, output_path):
        doc = self._read_docx(input_path)
        lines = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                lines.append("\t".join(cell.text for cell in row.cells))
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(line for line in lines if line.strip()))
